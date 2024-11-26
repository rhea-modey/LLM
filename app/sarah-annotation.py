import openai
import streamlit as st
from openai import OpenAI
import os
from docx import Document
import json
import random

# client = OpenAI()
openai.api_key = os.getenv("API_key")

# Set up Streamlit app
st.title("Automated Content Analysis with Chatbots")

# Define chatbot personas
chatbot_personas = {
    "Emily": "You are a 45-year-old Caucasian female social scientist with" +
        "a Ph.D. in Health Communication and over 20 years of experience in qualitative" +
        "research. You are known for a meticulous approach to analysis, focusing on" +
        "precision and consistency. As you analyze the data, ensure that each element is" +
        "carefully examined and categorized. Pay close attention to the details, and make" +
        "decisions based on thorough reasoning. Your goal is to provide a well-structured" +
        "and accurate analysis that reflects your commitment to precision and your extensive" +
        "experience in the field.",
    "Michael": "You are a 38-year-old Hispanic male social scientist with" +
        "a Ph.D. in Sociology and 15 years of experience in analyzing social dynamics and" +
        "health narratives. You are known for your intuitive and empathetic approach to" +
        "research, focusing on the emotional tone and social context. As you analyze the" +
        "data, consider the broader implications and the underlying human experiences. Your" +
        "goal is to capture the nuances and emotional depth of the data, reflecting your" +
        "understanding of the social dynamics and your commitment to empathy and insight."
}

# Phase 1: Chatbot persona creation
st.header("Phase 1: Chatbot Persona Creation")
persona_selected = st.radio("Select Chatbot Persona", list(chatbot_personas.keys()))
st.write(chatbot_personas[persona_selected])

# Phase 2: Codebook input
st.header("Phase 2: Codebook Input")
uploaded_codebook = st.file_uploader("Upload the Codebook (Word document)", type=["docx"])

if uploaded_codebook:
    doc = Document(uploaded_codebook)
    codebook = {}
    for para in doc.paragraphs:
        if para.text.strip():
            code, description = para.text.split(": ", 1)
            codebook[code.strip()] = description.strip()

    st.write("Codebook:", codebook)


# Phase 3: Independent Annotation
st.header("Phase 3: Independent Annotation")

# Sample text entries
text_entries = [
    "I started chemotherapy on February 10, 2020…After that I will have 25 days of radiation. Reconstruction will begin six months after that. So, 2020 has not been the year I hoped it would be. My ordeal combined with the COVID-19 pandemic has been surreal. But through it all, I have had great support from my family and friends.",
    "My name is Nikia. I was diagnosed with breast cancer at 16 years old in 1994 at a time when breast cancer treatment options were limited. Not only that – I was fighting for my life at a time when all of my friends’ biggest concerns were which dress they’d wear to prom. As you can imagine, breast cancer rocked my world. .",
    "Four kids and metastatic breast cancer. Tabatha Ann’s powerful story explains the realities of living with metastatic breast cancer while being a mom. It’s not easy but she refuses to ever give up.",
    "This is a picture of my best friend for the last 46 years, who has put his life on hold for a year to support me and stay by my side every day. That is what love is all about.” – Kathy, breast cancer survivor. Yesterday marked Kathy’s last day of treatment – join us in celebrating this incredible milestone with her!",
    "I was diagnosed with Stage 2A Invasive Ductal Carcinoma, at the age of 32, and since this day my life has completely changed",
    "Rest in peace to Jill Cohen, a powerful breast cancer advocate and friend, who passed away after 17 years of fighting breast cancer. Our hearts are with her family and friends.",
    "I had a bilateral mastectomy and had decided I did not want a reconstruction. It took a lot of work to feel at peace with my decision. My breasts had fed both of my children and served me well, but now it was time to let them go. I feel proud to still be here and to have highlighted that beauty comes in different shapes and sizes. It is what is inside us that shines out. Today, I am enough. Boobless and all.'",
    # Add more text entries as needed
]

annotations = {entry: {"Emily": [], "Michael": []} for entry in text_entries}

def annotate_based_on_persona(entry, persona):
    E_list = []
    M_list = []
    if persona == "Emily":
        # Emily is by the book
        if "prevent" in entry:
            E_list.append(codebook.get("1", "Prevention"))
        if "detect" in entry or "diagnos" in entry:
            E_list.append(codebook.get("2", "Detection, diagnosis"))
        if "chemo" in entry or "masectomy" in entry:
            E_list.append(codebook.get("3", "Treatment"))
        if "surviv" in entry:
            E_list.append(codebook.get("4", "Survivorship"))
        return E_list
    elif persona == "Michael":
        # Michael is more empathetic
        if "family history" in entry:
            M_list.append(codebook.get("1", "Prevention"))
        if "discover" in entry:
            M_list.append(codebook.get("2", "Detection, diagnosis"))
        if "treatment" in entry or "fight" in entry:
            M_list.append(codebook.get("3", "Treatment"))
        if "alive" in entry:
            M_list.append(codebook.get("4", "Survivorship"))
        return M_list

for entry in text_entries:
    st.subheader(f"Text Entry: {entry}")
    if persona_selected == "Emily":
        default_annotations = annotate_based_on_persona(entry, "Emily")
        annotations[entry]["Emily"] = st.multiselect(f"Emily's Annotation for: '{entry}'", list(codebook.values()), default=default_annotations, key=f"A_{entry}")
    elif persona_selected == "Michael":
        default_annotations = annotate_based_on_persona(entry, "Michael")
        annotations[entry]["Michael"] = st.multiselect(f"Michael's Annotation for: '{entry}'", list(codebook.values()), default=default_annotations, key=f"B_{entry}")

# Ensure that both chatbots can independently annotate text entries even if not selected in the persona creation phase
for entry in text_entries:
    if not annotations[entry]["Emily"]:
        annotations[entry]["Emily"] = annotate_based_on_persona(entry, "Emily")
    if not annotations[entry]["Michael"]:
        annotations[entry]["Michael"] = annotate_based_on_persona(entry, "Michael")

# Phase 4: Discuss Inconsistent Annotations
st.header("Phase 4: Discuss Inconsistent Annotations")

def simulate_discussion(entry, annotation_a, annotation_b):
    # This is a simple simulation. Replace with actual GPT-4o-mini logic.
    resolved_annotation = list(set(annotation_a + annotation_b))  # Combine both annotations and remove duplicates
    return resolved_annotation

updated_annotations = annotations.copy()

for entry, annotation in annotations.items():
    if set(annotation["Emily"]) != set(annotation["Michael"]):
        st.subheader(f"Discussion for: {entry}")
        discussion_result = simulate_discussion(entry, annotation["Emily"], annotation["Michael"])
        st.write(f"After discussion, we agreed that '{entry}' should be labeled as: {discussion_result}")
        updated_annotations[entry] = {"Emily": discussion_result, "Michael": discussion_result}
    else:
        updated_annotations[entry] = annotation

# Save updated codebook
if st.button("Save Updated Codebook"):
    updated_doc = Document()
    for code, description in codebook.items():
        updated_doc.add_paragraph(f"{code}: {description}")
    
    updated_doc.add_paragraph("\nUpdated Annotations:")
    for entry, annotations in updated_annotations.items():
        updated_doc.add_paragraph(f"{entry}: {', '.join(annotations['Emily'])}")

    updated_doc.save("updated_codebook.docx")
    with open("updated_codebook.docx", "rb") as f:
        st.download_button(label="Download Updated Codebook", data=f, file_name="updated_codebook.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Display updated codebook
st.header("Updated Codebook")
st.write(updated_annotations)
